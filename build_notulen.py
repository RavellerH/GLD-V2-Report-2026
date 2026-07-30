from __future__ import annotations

import html
import os
import sys
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
OUT = ROOT / "Deliverables"
OUT.mkdir(exist_ok=True)

SKILL_SCRIPTS = Path(
    r"C:\Users\HP\.codex\plugins\cache\openai-primary-runtime\documents"
    r"\26.709.11516\skills\documents\scripts"
)
sys.path.insert(0, str(SKILL_SCRIPTS))
from table_geometry import apply_table_geometry  # noqa: E402


NAVY = "17365D"
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "20252B"
MUTED = "667085"
LINE = "D7DEE8"
LIGHT = "F2F4F7"
PALE_BLUE = "E8EEF5"
PALE_GOLD = "FFF4CC"
WHITE = "FFFFFF"
GREEN = "1F6D4A"
AMBER = "8A5A00"
RED = "9B1C1C"


M24 = {
    "date_key": "24Jul2026",
    "date_display": "24 Juli 2026",
    "title": "NOTULEN MEETING",
    "subtitle": "Pengembangan Gas Leak Detector (GLD)",
    "meta": [
        ("Tanggal", "24 Juli 2026"),
        ("Waktu", "09.32–11.00 WIB"),
        ("Proyek", "Gas Leak Detector (GLD) — Real Testing & Certification"),
        (
            "Peserta",
            "PT LAPI Ganesha Utama (LGU): Pak Muhammad, Pak Farhan, Ilma, Fahmi, Beni\n"
            "PT Pertamina: Pak Sena, Pak Maman, Pak Adit, Mas Roni, Mas Indra (RU IV Cilacap)",
        ),
    ],
    "agenda": [
        "Pembaruan progres pengembangan Gas Leak Detector (GLD).",
        "Status sertifikasi perangkat.",
        "Persiapan pengujian lapangan di RU IV Cilacap.",
        "Tinjauan teknis sensor, jaringan, dan aplikasi manajemen.",
        "Koordinasi jadwal dan persyaratan instalasi.",
    ],
    "discussion": [
        (
            "1. Status Pengembangan dan Sertifikasi",
            [
                (
                    "Versi 24 V (catu daya)",
                    [
                        "Pengembangan telah selesai dan perangkat siap memasuki proses sertifikasi.",
                        "Administrasi sertifikasi telah disiapkan.",
                        "Proses sertifikasi direncanakan menggunakan jasa biro sertifikasi di Shanghai.",
                    ],
                ),
                (
                    "Versi baterai",
                    [
                        "Masih dalam tahap pengembangan.",
                        "Tantangan utama adalah konsumsi daya sensor MQ sekitar 10 W atau total arus sekitar 1 A untuk delapan sensor.",
                        "Target optimasi adalah konsumsi arus di bawah 100 mA dengan masa operasi minimal 30 hari.",
                        "Desain baru dengan baterai isi ulang diproyeksikan berpotensi mencapai masa operasi 6–8 bulan.",
                    ],
                ),
            ],
        ),
        (
            "2. Teknologi dan Fitur",
            [
                (
                    "Deteksi sensor",
                    [
                        "Sistem menggunakan fusi data dari delapan sensor MQ.",
                        "Gas yang disebutkan untuk dideteksi: H₂, LPG, metana, CO₂, dan clean air. Target minimum dinyatakan enam kelas; daftar sumber hanya menyebut lima kategori sehingga kelas keenam masih perlu dikonfirmasi.",
                        "Model CNN 1D dilaporkan memiliki akurasi 97,6%. Klaim ini perlu direkonsiliasi dengan hasil model lain dan diverifikasi sebelum dijadikan baseline resmi.",
                        "Setiap unit GLD perlu dilatih atau dikalibrasi untuk gas target yang akan dideteksi.",
                    ],
                ),
                (
                    "Jaringan LoRa mesh",
                    [
                        "Topologi komunikasi: GLD → Cluster Head (CH) → Gateway → Server.",
                        "Jaringan mesh mendukung self-healing dan automatic failover.",
                        "Pengujian telah mencapai tiga hop dengan latensi kurang dari 5 detik pada kondisi normal dan 5–15 detik pada kondisi ramai.",
                        "Cluster Head mengonsumsi sekitar 10 W dan dilaporkan stabil beroperasi selama 20 hari nonstop menggunakan panel surya.",
                    ],
                ),
                (
                    "Aplikasi manajemen",
                    [
                        "Tersedia perangkat bantu untuk konfigurasi Gateway, Cluster Head, dan sensor GLD.",
                        "Monitoring real-time meliputi status koneksi, kekuatan sinyal (RSSI/SNR), dan heat map deteksi gas.",
                        "Fitur yang dibahas mencakup rekalibrasi jarak jauh, notifikasi alarm, dan pencatatan data.",
                    ],
                ),
            ],
        ),
        (
            "3. Persiapan Pengujian Lapangan — RU IV Cilacap",
            [
                (
                    "Rencana kunjungan",
                    [
                        "Rencana awal: survei pada akhir Juli 2026 dan instalasi pada September 2026.",
                        "Strategi dapat diubah menjadi satu kunjungan yang menggabungkan survei dan instalasi apabila kondisi memungkinkan.",
                        "Survei mencakup pemetaan area, penentuan posisi sensor/CH/Gateway, dan analisis coverage.",
                    ],
                ),
                (
                    "Persyaratan administratif",
                    [
                        "Survei awal menggunakan izin visitor dengan durasi maksimal tiga hari.",
                        "Instalasi memerlukan izin kerja dan surat masuk barang.",
                        "Dokumen HSE dan APD mengikuti template RU IV Cilacap; APD warna merah perlu dihindari.",
                        "Koordinasi dengan tim HSE RU IV Cilacap wajib dilakukan sejak awal.",
                    ],
                ),
                (
                    "Rencana deployment",
                    [
                        "Dibahas dua sistem paralel, yaitu server lokal RU dan cloud, untuk perbandingan performa.",
                        "Perangkat portable untuk pemetaan sinyal LoRa telah disiapkan.",
                        "Test chamber portable akan dibawa untuk kalibrasi di lokasi.",
                    ],
                ),
            ],
        ),
        (
            "4. Pengembangan Tambahan",
            [
                (
                    "Deteksi api berbasis kamera",
                    [
                        "Model dilaporkan telah dikembangkan dengan akurasi 79% dan masih perlu ditingkatkan.",
                        "Perangkat yang digunakan adalah Jetson Nano, yang telah dibeli.",
                        "Dataset bersumber dari Google dan menggunakan kamera inframerah.",
                        "Kamera direncanakan dicari di China bersamaan dengan perjalanan sertifikasi.",
                    ],
                ),
                (
                    "Integrasi sensor arah angin",
                    [
                        "Sensor arah angin direncanakan sebagai bagian dari sistem peringatan dini.",
                        "Tujuannya adalah mendukung prediksi dispersi gas dan penyebaran api untuk respons darurat.",
                    ],
                ),
            ],
        ),
        (
            "5. Masukan dan Diskusi Teknis",
            [
                (
                    "",
                    [
                        "Pertamina meminta frekuensi pengiriman data disesuaikan dengan kebutuhan keselamatan.",
                        "Benchmark IoT monitoring yang dibahas adalah masa baterai lima tahun dengan transmisi dua kali per hari.",
                        "Sistem diharapkan plug-and-play dengan konfigurasi manual seminimal mungkin.",
                        "Konfigurasi default jaringan harus optimal; penyetelan manual hanya digunakan untuk kasus khusus.",
                    ],
                )
            ],
        ),
    ],
    "actions": [
        ["1", "Optimasi konsumsi daya versi baterai (<100 mA) dan uji masa operasi minimal 30 hari.", "Tim LGU (Pak Muhammad)", "15 Agu 2026", "Berjalan"],
        ["2", "Finalisasi desain baru versi baterai dengan target masa operasi 6–8 bulan.", "Tim LGU", "15 Agu 2026", "Berjalan"],
        ["3", "Sewa ulang tabung gas hidrogen (H₂) untuk proses pelatihan model GLD.", "Pak Farhan", "5 Agu 2026", "Tertunda"],
        ["4", "Verifikasi kemampuan deteksi H₂, LPG, metana, CO₂, clean air, dan konfirmasi kelas keenam.", "Pak Farhan, Ilma", "10 Agu 2026", "Berjalan"],
        ["5", "Uji multi-sensor simultan pada jaringan LoRa mesh (stress test).", "Beni, Fahmi", "10 Agu 2026", "Berjalan"],
        ["6", "Implementasikan retry/recovery untuk data alarm yang hilang saat transmisi.", "Beni", "15 Agu 2026", "Belum mulai"],
        ["7", "Unggah progress report dan timeline ke platform cloud serta bagikan akses kepada Pertamina.", "Pak Farhan", "31 Jul 2026", "Tertunda"],
        ["8", "Siapkan VPS kecil untuk deployment sistem monitoring berbasis cloud.", "Pak Muhammad", "20 Agu 2026", "Belum mulai"],
        ["9", "Sempurnakan aplikasi manajemen dan konfigurasi default agar plug-and-play.", "Fahmi", "20 Agu 2026", "Berjalan"],
        ["10", "Koordinasikan jadwal survei RU IV Cilacap dengan Pak Maman, Pak Adit, dan tim HSE.", "Pak Sena (Pertamina)", "31 Jul 2026", "Tertunda"],
        ["11", "Bagikan template izin kerja, surat masuk barang, dan dokumen HSE dari RU IV Cilacap.", "Pak Maman (Pertamina)", "31 Jul 2026", "Tertunda"],
        ["12", "Siapkan dokumen HSE dan perizinan masuk kilang, termasuk CS Negara.", "Tim LGU (Pak Muhammad)", "15 Agu 2026", "Belum mulai"],
        ["13", "Adakan APD warna non-merah untuk tim LGU.", "Pak Farhan", "10 Agu 2026", "Tertunda"],
        ["14", "Siapkan perangkat portable untuk pemetaan sinyal LoRa di lapangan.", "Pak Farhan", "20 Agu 2026", "Siap"],
        ["15", "Siapkan test chamber portable dan tabung gas untuk pengujian di lokasi.", "Tim LGU", "20 Agu 2026", "Berjalan"],
        ["16", "Aktif berkomunikasi di WhatsApp Group dengan tim Instrumen dan Merchant Insurance Pertamina.", "Semua pihak", "Berkelanjutan", "Berjalan"],
        ["17", "Siapkan perjalanan ke China untuk sertifikasi di Shanghai dan pencarian kamera di Wuhan.", "Pak Muhammad", "TBD — setelah versi baterai siap", "Ditahan"],
        ["18", "Laksanakan meeting persiapan pra-deployment bersama tim RU IV Cilacap (RAP meeting).", "Pak Sena, Pak Maman, Tim LGU", "15 Agu 2026", "Belum mulai"],
    ],
    "decisions": [
        "Deployment RU IV Cilacap akan menggunakan strategi dual system: server lokal dan cloud untuk evaluasi performa.",
        "Survei dan instalasi akan digabung dalam satu kunjungan apabila memungkinkan, dengan target saat itu September 2026.",
        "Kemampuan deteksi minimum ditetapkan enam kelas. H₂, LPG, metana, CO₂, dan clean air telah disebutkan; kelas keenam perlu dikonfirmasi.",
        "Pengembangan versi baterai dilanjutkan secara paralel dan tidak menghambat deployment versi 24 V.",
        "Koordinasi dengan tim HSE RU IV Cilacap wajib dimulai sejak awal untuk mencegah penolakan perizinan.",
    ],
    "risks": [
        ["Masa operasi baterai tidak mencapai target minimum 30 hari.", "Lanjutkan deployment versi 24 V; versi baterai menjadi fase pengembangan berikutnya."],
        ["Perizinan masuk kilang terlambat atau ditolak.", "Koordinasi awal dengan tim HSE/PIC RU IV Cilacap dan siapkan dokumen lengkap jauh hari."],
        ["Coverage LoRa mesh tidak optimal di lapangan.", "Lakukan survei dengan perangkat pemetaan sinyal dan tambahkan Cluster Head bila diperlukan."],
        ["Data alarm hilang saat transmisi.", "Implementasikan retry, local buffering, dan heartbeat untuk mendeteksi kehilangan koneksi."],
    ],
    "next": [
        ("Agenda", "RAP meeting persiapan deployment RU IV Cilacap."),
        ("Peserta", "Tim LGU dan tim RU IV Cilacap, termasuk HSE, Instrumen, dan Operasi."),
        ("Target", "Sekitar 15 Agustus 2026."),
    ],
    "notes": [
        "Komunikasi melalui WhatsApp Group perlu lebih aktif untuk mempercepat umpan balik teknis.",
        "Progress report perlu dapat diakses secara real-time oleh Pertamina.",
        "Setiap perubahan ruang lingkup atau timeline perlu segera dikomunikasikan.",
        "Klaim akurasi CNN 1D 97,6% dan flame detection 79% dicatat sebagai laporan meeting dan perlu verifikasi teknis lebih lanjut.",
    ],
}


M30 = {
    "date_key": "30Jul2026",
    "date_display": "30 Juli 2026",
    "title": "NOTULEN MEETING",
    "subtitle": "Gas Leak Detection (GLD) Tahap 2",
    "meta": [
        ("Tanggal", "30 Juli 2026"),
        ("Waktu", "Tidak tercantum pada catatan sumber"),
        ("Proyek", "Gas Leak Detection (GLD) — Weekly Progress & Technical Discussion"),
        (
            "Peserta",
            "Tim LGU / Lab IoT ITB (weekly internal). Daftar lengkap tidak tercantum. "
            "Catatan menyebut permintaan Pak Pindoan (Pertamina) dan tindak lanjut anggaran kepada Pak Tresnandi.",
        ),
    ],
    "agenda": [
        "Arsitektur aplikasi lokal/cloud dan monitoring terpusat.",
        "Optimasi daya serta duty-cycle GLD portable.",
        "Peran Wi-Fi dan LoRa dalam komunikasi data.",
        "Persiapan kunjungan lapangan RU IV Cilacap.",
        "Gas chamber, karakterisasi sensor MQ, dan status model AI.",
        "Casing, reliabilitas sistem, dokumentasi, konfigurasi, anggaran, dan logistik.",
    ],
    "discussion": [
        (
            "1. Arsitektur Aplikasi dan Monitoring",
            [
                (
                    "",
                    [
                        "Aplikasi dan database dibuat penuh secara lokal; cloud hanya berperan sebagai gateway/relay.",
                        "Instalasi direncanakan pada komputer site dan kantor Jakarta agar pusat dapat memonitor, sesuai permintaan Pak Pindoan.",
                        "Biaya cloud bila hanya sebagai gateway disebut USD 5, tetapi satuannya—per bulan atau per perangkat—belum dikonfirmasi.",
                    ],
                )
            ],
        ),
        (
            "2. Daya dan GLD Portable",
            [
                (
                    "",
                    [
                        "Target konsumsi di bawah 100 mA dengan on-off cepat memerlukan modifikasi hardware.",
                        "Push alarm belum diuji.",
                        "Versi portable/baterai memerlukan magnet untuk mounting; sumber pengadaan belum ditentukan.",
                        "Duty-cycle off berpotensi membuat perangkat melewatkan keberadaan gas. Diperlukan protokol missed-report untuk memeriksa baterai, sensor, dan kondisi perangkat bila laporan rutin tidak diterima.",
                    ],
                )
            ],
        ),
        (
            "3. Komunikasi: Wi-Fi dan LoRa",
            [
                (
                    "",
                    [
                        "Wi-Fi ESP dapat digunakan sebagai kanal konfigurasi atau serial nirkabel lokal dan memerlukan perintah enable/disable.",
                        "LoRa tetap selalu aktif sebagai jalur telemetri utama.",
                        "Antena Wi-Fi 2,4 GHz masih berada di dalam casing dan perlu dipindahkan ke luar.",
                    ],
                )
            ],
        ),
        (
            "4. Persiapan Kunjungan RU IV Cilacap",
            [
                (
                    "",
                    [
                        "Kunjungan dijadwalkan pada 9–10 Agustus 2026: berangkat Minggu, 9 Agustus, dan instalasi Senin, 10 Agustus.",
                        "Jadwal tim pada 10 Agustus perlu dikosongkan.",
                        "Rencana ini menggabungkan survei dan instalasi dalam satu kunjungan, menggantikan rencana 24 Juli yang memisahkan survei akhir Juli dan instalasi September 2026.",
                    ],
                )
            ],
        ),
        (
            "5. Gas Chamber dan Karakterisasi Sensor MQ",
            [
                (
                    "",
                    [
                        "Gas chamber perlu dibuat portable agar mudah dibawa ke site dan dapat digunakan untuk mencoba gas yang sulit tersedia di laboratorium.",
                        "Dimensi dan berat test chamber belum ditetapkan.",
                        "Sensor MQ perlu dipanaskan sampai stabil sebelum metode on-off diuji. Waktu pemanasan awal yang dibahas adalah sekitar 30 menit.",
                        "Informasi temperatur heater sekitar 200 °C berasal dari ChatGPT dan belum diverifikasi terhadap datasheet resmi.",
                        "Penyemprotan oksigen menyebabkan tegangan sensor turun (respons negatif).",
                        "Perlu dievaluasi pemanasan cepat atau penjagaan temperatur agar panas sensor tidak turun terlalu cepat.",
                        "Pengaruh duty-cycle terhadap temperatur kerja sensor berpotensi menjadi fitur tambahan untuk machine learning.",
                        "Sensor temperatur tambahan tidak dapat dipasang karena keterbatasan ruang. Temperatur lapangan perlu dicatat manual pada kunjungan Agustus.",
                        "LoRa telah diuji pada pengujian sebelumnya.",
                    ],
                )
            ],
        ),
        (
            "6. Status dan Kekhawatiran Model AI",
            [
                (
                    "",
                    [
                        "Saat ini baru tiga kelas yang tervalidasi: clean air, LPG, dan CO₂.",
                        "Uji penyemprotan LPG langsung dilaporkan berhasil.",
                        "Model yang di-deploy ke ESP32 telah dipangkas sekitar 84% dengan akurasi yang dilaporkan tetap terjaga.",
                        "Data sensitivitas sensor MQ terhadap gas telah ditambahkan ke dokumentasi.",
                        "Akurasi model yang sangat tinggi menimbulkan kekhawatiran; perlu verifikasi ulang terhadap potensi overfitting atau data leakage.",
                    ],
                )
            ],
        ),
        (
            "7. Casing dan Perlindungan Hardware",
            [
                (
                    "",
                    [
                        "Terdapat risiko genangan air di dalam casing akibat kondensasi.",
                        "Opsi mitigasi yang dibahas meliputi lubang sirkulasi kecil, kipas berkala, atau pemanas. Pilihan final belum ditetapkan.",
                    ],
                )
            ],
        ),
        (
            "8. Reliabilitas Sistem",
            [
                (
                    "",
                    [
                        "Potensi data collision pada komunikasi LoRa multi-sensor dinilai kecil, tetapi tetap perlu diuji secara eksplisit.",
                    ],
                )
            ],
        ),
        (
            "9. Dokumentasi dan Konfigurasi",
            [
                (
                    "",
                    [
                        "Progres perlu didokumentasikan sebagai knowledge base internal tim, bukan hanya untuk kebutuhan klien.",
                        "Parameter konfigurasi akan disederhanakan; operator tetap perlu memahami konfigurasi inti yang dipertahankan.",
                    ],
                )
            ],
        ),
        (
            "10. Anggaran dan Logistik",
            [
                (
                    "",
                    [
                        "Sisa anggaran perlu dikonfirmasi kepada Pak Tresnandi.",
                        "Laboratorium memiliki kabel HDMI panjang untuk sesi berikutnya.",
                        "Istilah “susi-sensor proposal” pada catatan sumber belum jelas dan perlu diklarifikasi bila relevan.",
                    ],
                )
            ],
        ),
    ],
    "actions": [
        ["1", "Modifikasi hardware untuk konsumsi <100 mA dengan on-off cepat.", "Belum ditetapkan", "Belum ditetapkan", "Terbuka"],
        ["2", "Uji push alarm.", "Belum ditetapkan", "Belum ditetapkan", "Belum diuji"],
        ["3", "Buat perintah enable/disable Wi-Fi untuk mode konfigurasi/serial lokal.", "Belum ditetapkan", "Belum ditetapkan", "Terbuka"],
        ["4", "Pindahkan antena Wi-Fi 2,4 GHz ke luar casing.", "Belum ditetapkan", "Belum ditetapkan", "Terbuka"],
        ["5", "Cari magnet untuk mounting GLD portable versi baterai.", "Belum ditetapkan", "Belum ditetapkan", "Terbuka"],
        ["6", "Uji potensi data collision pada LoRa multi-sensor.", "Belum ditetapkan", "Belum ditetapkan", "Perlu diuji"],
        ["7", "Definisikan protokol missed-report untuk memeriksa baterai/sensor bila laporan rutin tidak diterima.", "Belum ditetapkan", "Belum ditetapkan", "Terbuka"],
        ["8", "Sederhanakan parameter konfigurasi tanpa menghilangkan kontrol inti operator.", "Belum ditetapkan", "Belum ditetapkan", "Terbuka"],
        ["9", "Tentukan dimensi dan berat test chamber portable.", "Belum ditetapkan", "Belum ditetapkan", "Terbuka"],
        ["10", "Konfirmasi sisa anggaran kepada Pak Tresnandi.", "Belum ditetapkan", "Belum ditetapkan", "Terbuka"],
        ["11", "Verifikasi ulang akurasi model CNN 1D/TCN dan cek potensi data leakage/overfitting.", "Belum ditetapkan", "Belum ditetapkan", "Prioritas"],
        ["12", "Tentukan mekanisme anti-kondensasi casing.", "Belum ditetapkan", "Belum ditetapkan", "Opsi belum final"],
        ["13", "Catat temperatur lapangan untuk analisis duty-cycle dan temperatur sensor.", "Tim kunjungan", "9–10 Agu 2026", "Terjadwal"],
        ["14", "Instal aplikasi lokal dan database di site RU IV Cilacap serta kantor Jakarta.", "Belum ditetapkan", "Belum ditetapkan", "Permintaan Pak Pindoan"],
        ["15", "Konfirmasi satuan biaya cloud gateway-only sebesar USD 5.", "Belum ditetapkan", "Belum ditetapkan", "Terbuka"],
    ],
    "decisions": [
        "Aplikasi dan database penuh berjalan lokal di site RU IV Cilacap serta kantor Jakarta; cloud hanya sebagai gateway/relay.",
        "Wi-Fi digunakan untuk konfigurasi/serial lokal; LoRa tetap menjadi jalur telemetri utama dan selalu aktif.",
        "Survei dan instalasi RU IV Cilacap digabung pada 10 Agustus 2026, dengan keberangkatan 9 Agustus.",
        "Dokumentasi progres juga difungsikan sebagai knowledge base internal tim.",
        "Parameter konfigurasi akan disederhanakan dengan tetap mempertahankan konfigurasi inti yang perlu dipahami operator.",
    ],
    "risks": [
        ["Kondensasi menyebabkan genangan air di dalam casing.", "Evaluasi lubang sirkulasi kecil, kipas berkala, atau pemanas; opsi final belum dipilih."],
        ["Data collision pada LoRa multi-sensor.", "Lakukan pengujian eksplisit meskipun probabilitas dinilai kecil."],
        ["GLD portable melewatkan gas saat duty-cycle off.", "Definisikan missed-report dan pemeriksaan baterai/sensor bila laporan rutin tidak diterima."],
        ["Akurasi model terlalu tinggi atau tidak general.", "Verifikasi independen dan periksa overfitting serta data leakage."],
        ["Antena Wi-Fi 2,4 GHz berada di dalam casing.", "Pindahkan antena ke luar agar kanal konfigurasi lokal lebih andal."],
    ],
    "next": [
        ("Kegiatan", "Kunjungan lapangan, survei, dan instalasi RU IV Cilacap."),
        ("Keberangkatan", "Minggu, 9 Agustus 2026."),
        ("Pelaksanaan", "Senin, 10 Agustus 2026."),
    ],
    "notes": [
        "Daftar peserta lengkap, PIC, dan deadline sebagian besar action item tidak tercantum pada catatan sumber.",
        "Dimensi chamber, sisa anggaran, dan satuan biaya cloud masih merupakan pertanyaan terbuka.",
        "Temperatur heater sekitar 200 °C belum diverifikasi terhadap datasheet resmi.",
        "Istilah “susi-sensor proposal” diduga salah transkrip dan perlu klarifikasi bila relevan.",
    ],
}


def set_cell_margins(cell, top=90, start=120, bottom=90, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tag = OxmlElement("w:tblHeader")
    tag.set(qn("w:val"), "true")
    tr_pr.append(tag)


def set_run(run, size=10.5, color=INK, bold=False, italic=False, font="Calibri"):
    run.font.name = font
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), font)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), font)
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    run.bold = bold
    run.italic = italic


def set_repeat_safe(row):
    for cell in row.cells:
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_margins(cell)


def add_page_field(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r1 = paragraph.add_run("Halaman ")
    set_run(r1, size=8.5, color=MUTED)
    run = paragraph.add_run()
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    run._r.addnext(fld)


def create_numbering(doc):
    numbering = doc.part.numbering_part.element
    max_abs = max([int(x.get(qn("w:abstractNumId"))) for x in numbering.findall(qn("w:abstractNum"))] or [0])
    max_num = max([int(x.get(qn("w:numId"))) for x in numbering.findall(qn("w:num"))] or [0])

    def add_one(fmt, text, left, hanging, font=None):
        nonlocal max_abs, max_num
        max_abs += 1
        max_num += 1
        abstract = parse_xml(
            f'<w:abstractNum {nsdecls("w")} w:abstractNumId="{max_abs}">'
            f'<w:multiLevelType w:val="singleLevel"/>'
            f'<w:lvl w:ilvl="0"><w:start w:val="1"/><w:numFmt w:val="{fmt}"/>'
            f'<w:lvlText w:val="{text}"/><w:lvlJc w:val="left"/>'
            f'<w:pPr><w:tabs><w:tab w:val="num" w:pos="{left}"/></w:tabs>'
            f'<w:ind w:left="{left}" w:hanging="{hanging}"/></w:pPr>'
            + (f'<w:rPr><w:rFonts w:ascii="{font}" w:hAnsi="{font}"/></w:rPr>' if font else "")
            + "</w:lvl></w:abstractNum>"
        )
        numbering.append(abstract)
        num = parse_xml(
            f'<w:num {nsdecls("w")} w:numId="{max_num}">'
            f'<w:abstractNumId w:val="{max_abs}"/></w:num>'
        )
        numbering.append(num)
        return max_num

    return add_one("bullet", "•", 720, 360, "Arial"), add_one("decimal", "%1.", 720, 360)


def apply_num(paragraph, num_id):
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = p_pr.find(qn("w:numPr"))
    if num_pr is None:
        num_pr = OxmlElement("w:numPr")
        p_pr.append(num_pr)
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num_id_el = OxmlElement("w:numId")
    num_id_el.set(qn("w:val"), str(num_id))
    num_pr.extend([ilvl, num_id_el])


def add_bullet(doc, text, bullet_id, after=4):
    p = doc.add_paragraph()
    apply_num(p, bullet_id)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.10
    set_run(p.add_run(text), size=10.2)
    return p


def add_section_title(doc, text):
    p = doc.add_paragraph(style="Heading 1")
    p.paragraph_format.keep_with_next = True
    set_run(p.add_run(text), size=16, color=BLUE, bold=True)
    return p


def add_subheading(doc, text):
    p = doc.add_paragraph(style="Heading 2")
    p.paragraph_format.keep_with_next = True
    set_run(p.add_run(text), size=11.5, color=DARK_BLUE, bold=True)
    return p


def style_doc(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.78)
    section.bottom_margin = Inches(0.72)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.36)
    section.footer_distance = Inches(0.36)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    h1 = doc.styles["Heading 1"]
    h1.font.name = "Calibri"
    h1.font.size = Pt(16)
    h1.font.bold = True
    h1.font.color.rgb = RGBColor.from_string(BLUE)
    h1.paragraph_format.space_before = Pt(15)
    h1.paragraph_format.space_after = Pt(7)
    h1.paragraph_format.keep_with_next = True

    h2 = doc.styles["Heading 2"]
    h2.font.name = "Calibri"
    h2.font.size = Pt(11.5)
    h2.font.bold = True
    h2.font.color.rgb = RGBColor.from_string(DARK_BLUE)
    h2.paragraph_format.space_before = Pt(9)
    h2.paragraph_format.space_after = Pt(4)
    h2.paragraph_format.keep_with_next = True

    title = doc.styles["Title"]
    title.font.name = "Calibri"
    title._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    title._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    title.font.size = Pt(24)
    title.font.bold = True
    title.font.color.rgb = RGBColor.from_string(NAVY)
    title.paragraph_format.space_before = Pt(6)
    title.paragraph_format.space_after = Pt(1)

    subtitle = doc.styles["Subtitle"]
    subtitle.font.name = "Calibri"
    subtitle._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    subtitle._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    subtitle.font.size = Pt(13)
    subtitle.font.bold = True
    subtitle.font.italic = False
    subtitle.font.color.rgb = RGBColor.from_string(MUTED)
    subtitle.paragraph_format.space_before = Pt(0)
    subtitle.paragraph_format.space_after = Pt(12)

    header = section.header.paragraphs[0]
    header.text = "GLD TAHAP 2  |  NOTULEN MEETING"
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_run(header.runs[0], size=8.5, color=MUTED, bold=True)
    footer = section.footer.paragraphs[0]
    add_page_field(footer)


def add_title_block(doc, meeting):
    p = doc.add_paragraph(style="Title")
    set_run(p.add_run(meeting["title"]), size=24, color=NAVY, bold=True)

    p = doc.add_paragraph(style="Subtitle")
    set_run(p.add_run(meeting["subtitle"]), size=13, color=MUTED, bold=True)

    for label, value in meeting["meta"]:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(1.15)
        p.paragraph_format.first_line_indent = Inches(-1.15)
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.line_spacing = 1.05
        set_run(p.add_run(f"{label.upper()}: "), size=8.7, color=NAVY, bold=True)
        set_run(p.add_run(value), size=9.5)
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(0)


def add_data_table(doc, headers, rows, widths, font_size=8.4):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    table.style = "Table Grid"
    hdr = table.rows[0]
    set_repeat_table_header(hdr)
    for i, title in enumerate(headers):
        hdr.cells[i].text = ""
        shade_cell(hdr.cells[i], NAVY)
        p = hdr.cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_run(p.add_run(title), size=8.3, color=WHITE, bold=True)
    set_repeat_safe(hdr)

    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = ""
            p = cells[i].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            if i in (0, 3, 4) and len(headers) == 5:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_run(p.add_run(str(value)), size=font_size)
            if len(headers) == 5 and i == 4:
                val = str(value).lower()
                if any(k in val for k in ("siap", "terjadwal")):
                    shade_cell(cells[i], "E8F5EE")
                elif any(k in val for k in ("prioritas", "tertunda", "belum")):
                    shade_cell(cells[i], PALE_GOLD)
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cells[i], top=80, bottom=80)
    apply_table_geometry(table, widths, table_width_dxa=9360, indent_dxa=120)
    return table


def build_docx(meeting):
    doc = Document()
    style_doc(doc)
    bullet_id, _ = create_numbering(doc)
    add_title_block(doc, meeting)

    add_section_title(doc, "Agenda")
    for item in meeting["agenda"]:
        add_bullet(doc, item, bullet_id)

    add_section_title(doc, "Ringkasan Pembahasan")
    for title, groups in meeting["discussion"]:
        add_subheading(doc, title)
        for label, points in groups:
            if label:
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(2)
                p.paragraph_format.space_after = Pt(3)
                p.paragraph_format.keep_with_next = True
                set_run(p.add_run(label), size=10.3, color=NAVY, bold=True)
            for point in points:
                add_bullet(doc, point, bullet_id)

    add_section_title(doc, "Action Items / Daftar Tugas")
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(7)
    set_run(
        p.add_run(
            "Status, PIC, dan tenggat mengikuti catatan sumber. Entri “belum ditetapkan” menandai informasi yang perlu dilengkapi."
        ),
        size=9.2,
        color=MUTED,
        italic=True,
    )
    add_data_table(
        doc,
        ["No.", "Tugas", "PIC", "Tenggat", "Status"],
        meeting["actions"],
        [520, 4120, 1840, 1450, 1430],
        font_size=7.9,
    )

    add_section_title(doc, "Keputusan Meeting")
    for item in meeting["decisions"]:
        add_bullet(doc, item, bullet_id)

    add_section_title(doc, "Risiko dan Mitigasi")
    add_data_table(doc, ["Risiko", "Mitigasi"], meeting["risks"], [3900, 5460], font_size=9.0)

    add_section_title(doc, "Next Meeting / Kegiatan Berikutnya")
    for label, value in meeting["next"]:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(3)
        set_run(p.add_run(f"{label}: "), size=10.2, color=NAVY, bold=True)
        set_run(p.add_run(value), size=10.2)

    add_section_title(doc, "Catatan Penting")
    for item in meeting["notes"]:
        add_bullet(doc, item, bullet_id)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(0)
    set_run(
        p.add_run("Dokumen ini disusun dari catatan meeting yang tersedia tanpa menambahkan PIC, keputusan, atau tanggal yang tidak tercatat."),
        size=8.7,
        color=MUTED,
        italic=True,
    )

    path = OUT / f"Notulen_Meeting_GLD_{meeting['date_key']}.docx"
    doc.save(path)
    return path


def html_bullets(items):
    return "<ul>" + "".join(f"<li>{html.escape(x)}</li>" for x in items) + "</ul>"


def html_table(headers, rows, cls=""):
    head = "".join(f"<th>{html.escape(str(h))}</th>" for h in headers)
    body = []
    for row in rows:
        cells = "".join(f"<td>{html.escape(str(v))}</td>" for v in row)
        body.append(f"<tr>{cells}</tr>")
    return f'<div class="table-wrap"><table class="{cls}"><thead><tr>{head}</tr></thead><tbody>{"".join(body)}</tbody></table></div>'


def build_html(meeting):
    meta = "".join(
        f'<div class="meta-label">{html.escape(label)}</div><div class="meta-value">{html.escape(value).replace(chr(10), "<br>")}</div>'
        for label, value in meeting["meta"]
    )
    discussion = []
    for title, groups in meeting["discussion"]:
        chunks = [f"<h3>{html.escape(title)}</h3>"]
        for label, points in groups:
            if label:
                chunks.append(f"<h4>{html.escape(label)}</h4>")
            chunks.append(html_bullets(points))
        discussion.append("".join(chunks))
    next_items = "".join(
        f"<p><strong>{html.escape(label)}:</strong> {html.escape(value)}</p>" for label, value in meeting["next"]
    )
    doc = f"""<!doctype html>
<html lang="id">
<head>
  <meta charset="utf-8">
  <meta name="viewport" content="width=device-width, initial-scale=1">
  <title>Notulen Meeting GLD — {html.escape(meeting["date_display"])}</title>
  <style>
    :root {{ --navy:#17365d; --blue:#2e74b5; --ink:#20252b; --muted:#667085; --line:#d7dee8; --pale:#e8eef5; --paper:#fff; }}
    * {{ box-sizing:border-box; }}
    body {{ margin:0; background:#edf1f5; color:var(--ink); font:15px/1.52 Calibri, "Segoe UI", Arial, sans-serif; }}
    main {{ width:min(960px, calc(100% - 32px)); margin:28px auto; background:var(--paper); padding:54px 64px 62px; box-shadow:0 12px 36px rgba(23,54,93,.12); }}
    .kicker {{ color:var(--muted); font-size:12px; font-weight:700; letter-spacing:.12em; text-transform:uppercase; }}
    h1 {{ margin:7px 0 0; color:var(--navy); font-size:34px; line-height:1.05; letter-spacing:-.02em; }}
    .subtitle {{ margin:5px 0 22px; color:var(--muted); font-size:19px; font-weight:650; }}
    .meta {{ display:grid; grid-template-columns:150px 1fr; border:1px solid var(--line); margin-bottom:28px; }}
    .meta-label,.meta-value {{ padding:9px 12px; border-bottom:1px solid var(--line); }}
    .meta-label:nth-last-child(-n+2),.meta-value:nth-last-child(-n+2) {{ border-bottom:0; }}
    .meta-label {{ background:var(--pale); color:var(--navy); font-size:12px; font-weight:800; letter-spacing:.05em; text-transform:uppercase; }}
    h2 {{ margin:30px 0 10px; padding-bottom:5px; color:var(--blue); font-size:22px; border-bottom:1px solid var(--line); }}
    h3 {{ margin:21px 0 7px; color:var(--navy); font-size:17px; }}
    h4 {{ margin:13px 0 4px; color:#1f4d78; font-size:15px; }}
    ul {{ margin:4px 0 12px; padding-left:23px; }}
    li {{ margin:4px 0; }}
    p {{ margin:5px 0; }}
    .note {{ color:var(--muted); font-size:13px; font-style:italic; }}
    .table-wrap {{ overflow-x:auto; margin:9px 0 22px; border:1px solid var(--line); }}
    table {{ width:100%; border-collapse:collapse; font-size:13px; }}
    th {{ padding:9px 8px; background:var(--navy); color:white; text-align:left; font-size:12px; letter-spacing:.02em; }}
    td {{ padding:8px; border-top:1px solid var(--line); vertical-align:top; }}
    tbody tr:nth-child(even) {{ background:#f8fafc; }}
    .actions th:first-child,.actions td:first-child {{ width:48px; text-align:center; }}
    .actions th:nth-child(4),.actions td:nth-child(4),.actions th:nth-child(5),.actions td:nth-child(5) {{ white-space:nowrap; }}
    footer {{ margin-top:34px; padding-top:10px; border-top:1px solid var(--line); color:var(--muted); font-size:12px; font-style:italic; }}
    @media (max-width:700px) {{
      main {{ width:100%; margin:0; padding:30px 20px 42px; box-shadow:none; }}
      .meta {{ grid-template-columns:110px 1fr; }}
      h1 {{ font-size:29px; }}
    }}
    @media print {{
      @page {{ size:Letter; margin:16mm 17mm; }}
      body {{ background:white; font-size:10.5pt; }}
      main {{ width:auto; margin:0; padding:0; box-shadow:none; }}
      h2,h3,h4 {{ break-after:avoid; }}
      table {{ font-size:8.5pt; }}
      tr {{ break-inside:avoid; }}
      .table-wrap {{ overflow:visible; }}
      a {{ color:inherit; text-decoration:none; }}
    }}
  </style>
</head>
<body>
<main>
  <header>
    <div class="kicker">GLD Tahap 2 · Notulen Resmi</div>
    <h1>{html.escape(meeting["title"])}</h1>
    <div class="subtitle">{html.escape(meeting["subtitle"])}</div>
    <div class="meta">{meta}</div>
  </header>
  <section><h2>Agenda</h2>{html_bullets(meeting["agenda"])}</section>
  <section><h2>Ringkasan Pembahasan</h2>{"".join(discussion)}</section>
  <section>
    <h2>Action Items / Daftar Tugas</h2>
    <p class="note">Status, PIC, dan tenggat mengikuti catatan sumber. Entri “belum ditetapkan” menandai informasi yang perlu dilengkapi.</p>
    {html_table(["No.", "Tugas", "PIC", "Tenggat", "Status"], meeting["actions"], "actions")}
  </section>
  <section><h2>Keputusan Meeting</h2>{html_bullets(meeting["decisions"])}</section>
  <section><h2>Risiko dan Mitigasi</h2>{html_table(["Risiko", "Mitigasi"], meeting["risks"])}</section>
  <section><h2>Next Meeting / Kegiatan Berikutnya</h2>{next_items}</section>
  <section><h2>Catatan Penting</h2>{html_bullets(meeting["notes"])}</section>
  <footer>Dokumen ini disusun dari catatan meeting yang tersedia tanpa menambahkan PIC, keputusan, atau tanggal yang tidak tercatat.</footer>
</main>
</body>
</html>
"""
    path = OUT / f"Notulen_Meeting_GLD_{meeting['date_key']}.html"
    path.write_text(doc, encoding="utf-8")
    return path


def main():
    outputs = []
    for meeting in (M24, M30):
        outputs.extend([build_docx(meeting), build_html(meeting)])
    for path in outputs:
        print(path)


if __name__ == "__main__":
    main()
